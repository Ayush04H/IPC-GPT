import os
import google.generativeai as genai
from docx import Document

# Configure API key
genai.configure(api_key="AIzaSyC-xV9TwibkPS31mUFr7z8IhGcfTxakr4o")

# Create the model
generation_config = {
    "temperature": 1.05,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

# Automate generation for a range of IPC sections
def generate_and_save_ipc_sections(start, end, save_dir=None):
    # Set save directory or use current directory
    if save_dir:
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)  # Create directory if it doesn't exist
    else:
        save_dir = os.getcwd()  # Use current directory

    chat_session = model.start_chat(history=[])
    for section_number in range(start, end + 1):
        # Generate prompt
        prompt = f"Describe IPC section {section_number} in 3000 words strictly"
        
        # Generate response
        print(f"Generating content for IPC Section {section_number}...")
        response = chat_session.send_message(prompt)
        content = response.text
        
        # Save to Word document
        file_name = os.path.join(save_dir, f"IPC Section {section_number}.docx")
        document = Document()
        document.add_heading(f"IPC Section {section_number}", level=1)
        document.add_paragraph(content)
        document.save(file_name)
        print(f"Saved content for IPC Section {section_number} to {file_name}")

# Call the function with a desired range and directory
generate_and_save_ipc_sections(67,77, save_dir="Sections")  # Change range and directory as needed
