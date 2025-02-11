import os
import google.generativeai as genai
from docx import Document

# Configure API key
genai.configure(api_key="")
# Create the model
generation_config = {
    "temperature": 1.05,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

# Automate generation for a list of IPC sections with descriptions
def generate_and_save_ipc_descriptions(sections, save_dir=None):
    """
    Generates descriptions for a list of IPC sections, including their names and statements.

    Args:
        sections (list): A list of dictionaries, where each dictionary contains
                       the section number and the section title.  Example:
                       [{'number': 1, 'title': 'Title and extent of operation of the Code.'},
                        {'number': 2, 'title': 'Punishment of offences committed within India.'}]
        save_dir (str, optional): The directory to save the generated documents.
                                   Defaults to the current working directory.
    """

    # Set save directory or use current directory
    if save_dir:
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)  # Create directory if it doesn't exist
    else:
        save_dir = os.getcwd()  # Use current directory

    chat_session = model.start_chat(history=[])
    for section in sections:
        section_number = section['number']
        section_title = section['title']

        # Generate prompt
        prompt = f"Describe IPC section {section_number}: '{section_title}' according to Indian Penal Code in 1000 words strictly"

        # Generate response
        print(f"Generating content for IPC Section {section_number}: {section_title}...")
        response = chat_session.send_message(prompt)
        content = response.text

        # Sanitize the section title for use in the filename
        safe_section_title = "".join(c if c.isalnum() or c in (' ', '_', '-') else "_" for c in section_title)
        safe_section_title = safe_section_title[:50]  # Truncate for filename length

        # Save to Word document
        file_name = os.path.join(save_dir, f"IPC Section {section_number} - {safe_section_title}.docx")
        document = Document()
        document.add_heading(f"IPC Section {section_number}: {section_title}", level=1)
        document.add_paragraph(content)
        document.save(file_name)
        print(f"Saved content for IPC Section {section_number}: {section_title} to {file_name}")


# Define the list of sections to process

sections_to_generate = [
    {'number': 289, 'title': 'Negligent conduct with respect to animal.'},
    {'number': 290, 'title': 'Punishment for public nuisance in cases not otherwise provided for.'},
    {'number': 291, 'title': 'Continuance of nuisance after injunction to discontinue.'},
    {'number': 292, 'title': 'Sale, etc., of obscene books, etc.'},
    {'number': 293, 'title': 'Sale, etc., of obscene objects to young person.'},
    {'number': 294, 'title': 'Obscene acts and songs.'},
    {'number': '294A', 'title': 'Keeping lottery office.'}
]

'''
sections_to_generate = [
    {'number': 295, 'title': 'Injuring or defiling place of work ship, with intent to insult the religion of any class.'},
    {'number': '295A', 'title': 'Deliberate and malicious acts, intended to outrage religious feelings of any class by insulting its religion or religious beliefs.'},
    {'number': 296, 'title': 'Disturbing religious assembly.'},
    {'number': 297, 'title': 'Trespassing on burial places, etc.'},
    {'number': 298, 'title': 'Uttering words, etc., with deliberate intent to wound the religious feelings.'},
    {'number': 299, 'title': 'Culpable homicide.'},
    {'number': 300, 'title': 'Murder. When culpable homicide is not murder.'},
    {'number': 301, 'title': 'Culpable homicide by causing death of person other than person whose death was intended.'},
    {'number': 302, 'title': 'Punishment for murder.'},
    {'number': 303, 'title': 'Punishment for murder by life-convict.'},
    {'number': 304, 'title': 'Punishment for culpable homicide not amounting to murder.'},
    {'number': '304A', 'title': 'Causing death by negligence.'},
    {'number': '304B', 'title': 'Dowry death.'},
    {'number': 305, 'title': 'Abetment of suicide of child or insane person.'},
    {'number': 306, 'title': 'Abetment of suicide.'},
    {'number': 307, 'title': 'Attempt to murder. Attempts by life-convicts.'},
    {'number': 308, 'title': 'Attempt to commit culpable homicide.'},
    {'number': 309, 'title': 'Attempt to commit suicide.'},
    {'number': 310, 'title': 'Thug.'},
    {'number': 311, 'title': 'Punishment.'},
    {'number': 312, 'title': 'Causing miscarriage.'},
    {'number': 313, 'title': 'Causing miscarriage without woman\'s consent.'},
    {'number': 314, 'title': 'Death caused by act done with intent to cause miscarriage. if act done without woman\'s consent.'},
    {'number': 315, 'title': 'Act done with intent to prevent child being born alive or to cause it to die after birth.'},
    {'number': 316, 'title': 'Causing death of quick unborn child by act amounting to culpable homicide.'},
    {'number': 317, 'title': 'Exposure and abandonment of child under twelve years, by parent or person having care of it.'},
    {'number': 318, 'title': 'Concealment of birth by secret disposal of dead body.'},
    {'number': 319, 'title': 'Hurt.'},
    {'number': 320, 'title': 'Grievous hurt.'},
    {'number': 321, 'title': 'Voluntarily causing hurt.'},
    {'number': 322, 'title': 'Voluntarily causing grievous hurt.'},
    {'number': 323, 'title': 'Punishment for voluntarily causing hurt.'},
    {'number': 324, 'title': 'Voluntarily causing hurt by dangerous weapons or means.'},
    {'number': 325, 'title': 'Punishment for voluntarily causing grievous hurt.'},
    {'number': 326, 'title': 'Voluntarily causing grievous hurt by dangerous weapons or means.'},
    {'number': '326A', 'title': 'Voluntarily causing grievous hurt by use of acid, etc.'},
    {'number': '326B', 'title': 'Voluntarily throwing or attempting to throw acid.'},
    {'number': 327, 'title': 'Voluntarily causing hurt to extort property, or to constrain to an illegal act.'},
    {'number': 328, 'title': 'Causing hurt by means of poison, etc., with intent to commit an offence.'},
    {'number': 329, 'title': 'Voluntarily causing grievous hurt to extort property, or to constrain to an illegal act.'},
    {'number': 330, 'title': 'Voluntarily causing hurt to extort confession, or to compel restoration of property.'},
    {'number': 331, 'title': 'Voluntarily causing grievous hurt to extort confession, or to compel restoration of property.'},
    {'number': 332, 'title': 'Voluntarily causing hurt to deter public servant from his duty.'},
    {'number': 333, 'title': 'Voluntarily causing grievous hurt to deter public servant from his duty.'},
    {'number': 334, 'title': 'Voluntarily causing hurt on provocation.'},
    {'number': 335, 'title': 'Voluntarily causing grievous hurt on provocation.'},
    {'number': 336, 'title': 'Act endangering life or personal safety of others.'},
    {'number': 337, 'title': 'Causing hurt by act endangering life or personal safety of others.'},
    {'number': 338, 'title': 'Causing grievous hurt by act endangering life or personal safety of others.'},
    {'number': 339, 'title': 'Wrongful restraint.'},
    {'number': 340, 'title': 'Wrongful confinement.'},
    {'number': 341, 'title': 'Punishment for wrongful restraint.'},
    {'number': 342, 'title': 'Punishment for wrongful confinement.'},
    {'number': 343, 'title': 'Wrongful confinement for three or more days.'},
    {'number': 344, 'title': 'Wrongful confinement for ten or more days.'},
    {'number': 345, 'title': 'Wrongful confinement of person for whose liberation writ has been issued.'},
    {'number': 346, 'title': 'Wrongful confinement in secret.'},
    {'number': 347, 'title': 'Wrongful confinement to extort property, or constrain to illegal act.'},
    {'number': 348, 'title': 'Wrongful confinement to extort confession, or compel restoration of property.'},
    {'number': 349, 'title': 'Force.'},
    {'number': 350, 'title': 'Criminal force.'},
    {'number': 351, 'title': 'Assault.'},
    {'number': 352, 'title': 'Punishment for assault or criminal force otherwise than on grave provocation.'},
    {'number': 353, 'title': 'Assault or criminal force to deter public servant from discharge of his duty.'},
    {'number': 354, 'title': 'Assault of criminal force to woman with intent to outrage her modesty.'},
    {'number': '354A', 'title': 'Sexual harassment and punishment for sexual harassment.'},
    {'number': '354B', 'title': 'Assault or use of criminal force to woman with intent to disrobe.'},
    {'number': '354C', 'title': 'Voyeurism.'},
    {'number': '354D', 'title': 'Stalking.'},
    {'number': 355, 'title': 'Assault or criminal force with intent to dishonour person, otherwise than on grave provocation.'},
    {'number': 356, 'title': 'Assault or criminal force in attempt to commit theft of property carried by a person.'},
    {'number': 357, 'title': 'Assault or criminal force in attempt wrongfully to confine a person.'},
    {'number': 358, 'title': 'Assault or criminal force on grave provocation.'}
]

sections_to_generate = [
    {'number': 360, 'title': 'Kidnapping from India.'},
    {'number': 361, 'title': 'Kidnapping from lawful guardianship.'},
    {'number': 362, 'title': 'Abduction.'},
    {'number': 363, 'title': 'Punishment for kidnapping.'},
    {'number': '363A', 'title': 'Kidnapping or maiming a minor for purposes of begging.'},
    {'number': 364, 'title': 'Kidnapping or abducting in order to murder.'},
    {'number': '364A', 'title': 'Kidnapping for ransom, etc.'},
    {'number': 365, 'title': 'Kidnapping or abducting with intent secretly and wrongfully to confine person.'},
    {'number': 366, 'title': 'Kidnapping, abducting or inducing woman to compel her marriage, etc.'},
    {'number': '366A', 'title': 'Procuration of minor girl.'},
    {'number': '366B', 'title': 'Importation of girl from foreign country.'},
    {'number': 367, 'title': 'Kidnapping or abducting in order to subject person to grievous hurt, slavery, etc.'},
    {'number': 368, 'title': 'Wrongfully concealing or keeping in confinement, kidnapped or abducted person.'},
    {'number': 369, 'title': 'Kidnapping or abducting child under ten years with intent to steal from its person.'},
    {'number': 370, 'title': 'Trafficking of person.'},
    {'number': '370A', 'title': 'Exploitation of a trafficked person.'},
    {'number': 371, 'title': 'Habitual dealing in slaves.'},
    {'number': 372, 'title': 'Selling minor for purposes of prostitution, etc.'},
    {'number': 373, 'title': 'Buying minor for purposes of prostitution, etc.'},
    {'number': 374, 'title': 'Unlawful compulsory labour.'},
    {'number': 375, 'title': 'Rape.'},
    {'number': 376, 'title': 'Punishment for rape.'},
    {'number': '376A', 'title': 'Punishment for causing death or resulting in persistent vegetative state of victim.'},
    {'number': '376B', 'title': 'Sexual intercourse by husband upon his wife during separation.'},
    {'number': '376C', 'title': 'Sexual intercourse by a person in authority.'},
    {'number': '376D', 'title': 'Gang rape.'},
    {'number': '376E', 'title': 'Punishment for repeat offenders.'},
    {'number': 378, 'title': 'Theft.'},
    {'number': 379, 'title': 'Punishment for theft.'},
    {'number': 380, 'title': 'Theft in dwelling house, etc.'},
    {'number': 381, 'title': 'Theft by clerk or servant of property in possession of master.'},
    {'number': 382, 'title': 'Theft after preparation made for causing death, hurt or restraint in order to the committing of the theft.'},
    {'number': 383, 'title': 'Extortion.'},
    {'number': 384, 'title': 'Punishment for extortion.'},
    {'number': 385, 'title': 'Putting person in fear of injury in order to commit extortion.'},
    {'number': 386, 'title': 'Extortion by putting a person in fear of death on grievous hurt.'},
    {'number': 387, 'title': 'Putting person in fear of death or of grievous hurt, in order to commit extortion.'},
    {'number': 388, 'title': 'Extortion by threat of accusation of an offence punishable with death or imprisonment for life, etc.'},
    {'number': 389, 'title': 'Putting person in fear of accusation of offence, in order to commit extortion.'},
    {'number': 390, 'title': 'Robbery. When theft is robbery. When extortion is robbery.'},
    {'number': 391, 'title': 'Dacoity.'},
    {'number': 392, 'title': 'Punishment for robbery.'},
    {'number': 393, 'title': 'Attempt to commit robbery.'},
    {'number': 394, 'title': 'Voluntarily causing hurt in committing robbery.'},
    {'number': 395, 'title': 'Punishment for dacoity.'},
    {'number': 396, 'title': 'Dacoity with murder.'},
    {'number': 397, 'title': 'Robbery, or dacoity, with attempt to cause death or grievous hurt.'},
    {'number': 398, 'title': 'Attempt to commit robbery or dacoity when armed with deadly weapon.'},
    {'number': 399, 'title': 'Making preparation to commit dacoity.'},
    {'number': 400, 'title': 'Punishment for belonging to gang of dacoits.'},
    {'number': 401, 'title': 'Punishment for belonging to gang of thieves.'},
    {'number': 402, 'title': 'Assembling for purpose of committing dacoity.'},
    {'number': 403, 'title': 'Dishonest misappropriation of property.'},
    {'number': 404, 'title': 'Dishonest misappropriation of property possessed by deceased person at the time of his death.'},
    {'number': 405, 'title': 'Criminal breach of trust.'},
    {'number': 406, 'title': 'Punishment for criminal breach of trust.'},
    {'number': 407, 'title': 'Criminal breach of trust by carrier, etc.'},
    {'number': 408, 'title': 'Criminal breach of trust by clerk or servant.'},
    {'number': 409, 'title': 'Criminal breach of trust by public, servant. or by banker, merchant or agent.'},
    {'number': 410, 'title': 'Stolen property.'},
    {'number': 411, 'title': 'Dishonestly receiving stolen property.'},
    {'number': 412, 'title': 'Dishonestly receiving property stolen in the commission of a dacoity.'},
    {'number': 413, 'title': 'Habitually dealing in stolen property.'},
    {'number': 414, 'title': 'Assisting in concealment of stolen property.'},
    {'number': 415, 'title': 'Cheating.'},
    {'number': 416, 'title': 'Cheating by personation.'},
    {'number': 417, 'title': 'Punishment for cheating.'},
    {'number': 418, 'title': 'Cheating with knowledge that wrongful loss may ensue to person whose interest offender is bound to protect.'},
    {'number': 419, 'title': 'Punishment for cheating by personation.'},
    {'number': 420, 'title': 'Cheating and dishonestly inducing delivery of property.'},
    {'number': 421, 'title': 'Dishonest or fraudulent removal or concealment of property to prevent distribution among creditor.'},
    {'number': 422, 'title': 'Dishonestly or fraudulently preventing debt being available for creditors.'},
    {'number': 423, 'title': 'Dishonest or fraudulent execution of deed of transfer containing false statement of consideration.'},
    {'number': 424, 'title': 'Dishonest or fraudulent removal or concealment of property.'},
    {'number': 425, 'title': 'Mischief.'},
    {'number': 426, 'title': 'Punishment for mischief.'},
    {'number': 427, 'title': 'Mischief causing damage to the amount of fifty rupees.'},
    {'number': 428, 'title': 'Mischief by killing or maiming animal of the value of ten rupees.'},
    {'number': 429, 'title': 'Mischief by killing or maiming cattle, etc., of any value or any animal of the value of fifty rupees.'},
    {'number': 430, 'title': 'Mischief by injury to works of irrigation or by wrongfully diverting water.'},
    {'number': 431, 'title': 'Mischief by injury to public road, bridge, river or channel.'},
    {'number': 432, 'title': 'Mischief by causing inundation or obstruction to public drainage attended with damage.'},
    {'number': 433, 'title': 'Mischief by destroying, moving or rendering less useful a light-house or sea-mark.'},
    {'number': 434, 'title': 'Mischief by destroying or moving, etc., a land-mark fixed by public authority.'},
    {'number': 435, 'title': 'Mischief by fire or explosive substance with intent to cause damage to amount of one hundred or (in case of agricultural produce ) ten rupees.'},
    {'number': 436, 'title': 'Mischief by fire or explosive substance with intent to destroy house, etc.'},
    {'number': 437, 'title': 'Mischief with intent to destroy or make unsafe a decked vessel or one of twenty tons burden.'},
    {'number': 438, 'title': 'Punishment for the mischief described in section 437 committed by fire or explosive substance.'},
    {'number': 439, 'title': 'Punishment for intentionally running vessel agroun, or ashore with intent to commit theft, etc.'},
    {'number': 440, 'title': 'Mischief committed after preparation made for causing death or hurt.'}
]
'''
# You can now print the list to verify
# for section in sections_to_generate:
#     print(f"Section {section['number']}: {section['title']}")
# Call the function with the list of sections and a save directory
generate_and_save_ipc_descriptions(sections_to_generate, save_dir="Sections_2")